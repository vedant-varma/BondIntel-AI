from tkinter import Tk, Label, Button, Entry, Listbox, END, ACTIVE, RAISED
from openpyxl import Workbook, load_workbook
from bond import Bond
from openAI import bond_predict, chat
from docx import Document
import os

# global variables
window = Tk()
bonds = []
global current_bond
doc = Document()


# extracts data from the worksheet and populates the bonds list
def extract_data(ws):
    for row in ws.iter_rows(min_row=2, values_only=True):
        bond = Bond(*row)
        bonds.append(bond)


# handles the click of the enter button
def enter_click():
    file_path = enter_box.get()
    enter_box.delete(0, END)
    wb = load_workbook(file_path)
    ws = wb.active
    extract_data(ws)
    for i, bond in enumerate(bonds):
        bond_list.insert(i, bond.name)


# handles the click of the show analysis button
def display_data():
    global current_bond
    bond_name = bond_list.get(bond_list.curselection())
    for bond in bonds:
        if bond.name == bond_name:
            current_bond = bond
            info_area.config(text=str(bond))


def ask_ai():
    global current_bond
    question = chat_box.get()
    chat_box.delete(0, END)
    answer = chat(current_bond, question)
    doc = Document('analysis.docx')
    doc.add_paragraph(f'Question: {question}')
    doc.add_paragraph(f'Answer: {answer}')
    doc.save('analysis.docx')
    if os.name == 'nt':  # for Windows
        os.startfile('analysis.docx')
    else:  # for MacOS/Linux
        os.system('open ' + 'analysis.docx')

# handles the click of the AI analysis button
def create_anls():
    bond_name = bond_list.get(bond_list.curselection())
    for bond in bonds:
        if bond.name == bond_name:
            p = doc.add_paragraph(bond_predict(bond))
            doc.save('analysis.docx')
            if os.name == 'nt':  # for Windows
                os.startfile('analysis.docx')
            else:  # for MacOS/Linux
                os.system('open ' + 'analysis.docx')


# configuring the window
window.geometry("1500x750")
window.title("test")
window.config(background="#040012")

# labels, buttons, and entry fields
intro_label = Label(window,
                    text="Enter Bond Data Below",
                    font=('Arial', 20, 'bold'),
                    fg='#045926',
                    relief=RAISED,
                    bd=10,
                    padx=10,
                    pady=10,
                    compound='top')
intro_label.place(x=50, y=100)

bond_enter_button = Button(window,
                           text="Submit",
                           command=enter_click,
                           font=("Arial", 20, 'bold'),
                           fg='#045926',
                           relief=RAISED,
                           bd=10,
                           padx=4,
                           pady=4,
                           compound='top',
                           state=ACTIVE)
bond_enter_button.place(x=400, y=150)

enter_box = Entry(window,
                  font=("Arial", 20, "bold"),
                  fg='#045926',
                  relief=RAISED,
                  bd=10)
enter_box.place(x=50, y=200)

bond_list = Listbox(window,
                    font=("Arial", 20, "bold"),
                    fg='#045926',
                    relief=RAISED,
                    bd=10)
bond_list.place(x=1150, y=100)

info_area = Label(window,
                  font=("Arial", 20, 'bold'),
                  fg='#045926',
                  relief=RAISED,
                  bd=10,
                  width=60,
                  height=11)
info_area.place(x=50, y=350)

show_analysis_btn = Button(window,
                           text="Bond Details",
                           command=display_data,
                           font=("Arial", 20, "bold"),
                           fg="#045926",
                           relief=RAISED,
                           bd=10,
                           padx=4,
                           pady=4,
                           state=ACTIVE)
show_analysis_btn.place(x=600, y=150)

create_anls_btn = Button(window,
                         text="Create Analysis",
                         command=create_anls,
                         font=("Arial", 20, "bold"),
                         fg="#045926",
                         relief=RAISED,
                         bd=10,
                         padx=4,
                         pady=4,
                         state=ACTIVE)
create_anls_btn.place(x=1150, y=460)

chat_box = Entry(window,
                 font=("Arial", 20, "bold"),
                 fg='#045926',
                 relief=RAISED,
                 bd=10)
chat_box.place(x=1150, y=600)

ask_ai_btn = Button(window,
                    text="Ask AI Expert",
                    command=ask_ai,
                    font=("Arial", 20, "bold"),
                    fg="#045926",
                    relief=RAISED,
                    bd=10,
                    padx=4,
                    pady=4,
                    state=ACTIVE)
ask_ai_btn.place(x=1150, y=660)

def run_gui():
    # start the GUI
    window.mainloop()
